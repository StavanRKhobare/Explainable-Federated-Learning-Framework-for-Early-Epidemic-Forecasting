import os
from docx import Document

def generate_txt(file_path, content):
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"[+] Generated: {file_path}")

def generate_docx(file_path, paragraphs):
    doc = Document()
    doc.add_heading("PATIENT MEDICAL RECORD", level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(file_path)
    print(f"[+] Generated: {file_path}")

def main():
    out_dir = "ehr_samples"
    os.makedirs(out_dir, exist_ok=True)

    # 1. Txt Positive
    txt_pos = """
    ========================================
    METROPOLITAN HOSPITAL - CLINICAL RECORD
    ========================================
    Patient Name: Aarav Sharma
    Age / Sex: 34 / Male
    Date: 2026-06-20
    
    SYMPTOMS & VITALS:
    Subject reports severe joint pain ("breakbone fever"), headaches behind eyes, and body rash.
    Recorded Vitals:
      - Blood Pressure: 120/80 mmHg
      - Pulse: 85 bpm
      - Temperature: 103.4 F (High Fever)
      
    DIAGNOSIS & LAB RESULTS:
    Dengue NS1 Antigen Test: POSITIVE (detected)
    IgM Antibody: Positive
    
    RECOMMENDATION:
    Prescribed paracetamol, strict bed rest, and aggressive hydration. Avoid aspirin/ibuprofen.
    """
    generate_txt(os.path.join(out_dir, "patient_dengue_pos_103F.txt"), txt_pos)

    # 2. Txt Negative
    txt_neg = """
    ========================================
    METROPOLITAN HOSPITAL - CLINICAL RECORD
    ========================================
    Patient Name: Priya Patel
    Age / Sex: 28 / Female
    Date: 2026-06-20
    
    SYMPTOMS & VITALS:
    Subject reports mild dry cough and slight runny nose. No body aches or rashes.
    Recorded Vitals:
      - Blood Pressure: 115/75 mmHg
      - Pulse: 72 bpm
      - Temperature: 98.6 F (Normal)
      
    DIAGNOSIS & LAB RESULTS:
    Dengue NS1 Antigen Test: NEGATIVE (not detected)
    COVID-19 Rapid Test: Negative
    
    RECOMMENDATION:
    Mild seasonal cold. Advised warm fluids and throat lozenges.
    """
    generate_txt(os.path.join(out_dir, "patient_dengue_neg_98F.txt"), txt_neg)

    # 3. Docx Positive
    docx_pos_paragraphs = [
        "Patient Name: Rajesh Kumar",
        "Age / Sex: 45 / Male",
        "Date of Examination: June 19, 2026",
        "Clinical Presentation: High fever of acute onset, muscle and joint pains (severe). Rash observed on upper limbs.",
        "Body Temperature: 39.5 C",
        "Laboratory Testing: Blood draw analyzed.",
        "Dengue Virus NS1 Rapid Test: Reactive (POSITIVE)",
        "Dengue IgG: Non-reactive",
        "Assessment: Early acute phase Dengue Infection. Follow-up platelet count scheduled in 48 hours."
    ]
    generate_docx(os.path.join(out_dir, "patient_rajesh_dengue_pos_39C.docx"), docx_pos_paragraphs)

    # 4. Docx Negative
    docx_neg_paragraphs = [
        "Patient Name: Sneha Reddy",
        "Age / Sex: 22 / Female",
        "Date of Examination: June 18, 2026",
        "Clinical Presentation: Patient presents with headache and mild fatigue. Denies joint pain or rash.",
        "Vitals: BP 110/70, Heart Rate 78, Temperature: 37.2 C",
        "Laboratory Diagnostics: Clinical testing conducted.",
        "Dengue NS1 Antigen: Negative (not detected)",
        "IgM Antibody: Non-reactive (negative)",
        "Assessment: Fatigue secondary to dehydration/overwork. Prescribed oral rehydration salts and rest."
    ]
    generate_docx(os.path.join(out_dir, "patient_sneha_dengue_neg_37C.docx"), docx_neg_paragraphs)

if __name__ == "__main__":
    main()
